"""
Document Loader Module for Multimodel Chatbot

This module handles document processing and preparation for the RAG system.
It provides:
- Loading documents from various file formats (PDF, TXT, DOCX)
- Text extraction and cleaning
- Document chunking into smaller pieces
- Metadata extraction and attachment
- Support for batch processing

The document loader prepares raw files into properly formatted chunks
that can be stored in the vector database.
"""

import os
import logging
from typing import List, Dict, Optional, Tuple
from pathlib import Path
from abc import ABC, abstractmethod

# Third-party imports
import PyPDF2
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain.schema import Document

# ============================================================================
# CONFIGURATION
# ============================================================================

logger = logging.getLogger(__name__)

# Supported file extensions
SUPPORTED_FORMATS = {
    '.pdf': 'PDF Document',
    '.txt': 'Text File',
    '.md': 'Markdown File',
    '.docx': 'Word Document',
}

# Maximum file size (100 MB)
MAX_FILE_SIZE = 100 * 1024 * 1024


# ============================================================================
# BASE FILE LOADER CLASS
# ============================================================================

class BaseFileLoader(ABC):
    """
    Abstract base class for file loaders.
    
    Defines the interface that all file type loaders must implement.
    """
    
    @abstractmethod
    def load(self, file_path: str) -> str:
        """Load and extract text from file."""
        pass
    
    @abstractmethod
    def supports_format(self, file_path: str) -> bool:
        """Check if loader supports this file format."""
        pass


# ============================================================================
# PDF LOADER
# ============================================================================

class PDFLoader(BaseFileLoader):
    """
    Loader for PDF documents.
    
    Extracts text from PDF files, handling multi-page documents.
    
    Attributes:
        file_path: Path to PDF file
    """
    
    def supports_format(self, file_path: str) -> bool:
        """Check if file is PDF."""
        return file_path.lower().endswith('.pdf')
    
    def load(self, file_path: str) -> str:
        """
        Extract text from PDF file.
        
        Args:
            file_path: Path to PDF file
            
        Returns:
            Extracted text from all pages
            
        Raises:
            FileNotFoundError: If file doesn't exist
            ValueError: If PDF is corrupted or unreadable
            
        Example:
            >>> loader = PDFLoader()
            >>> text = loader.load("document.pdf")
            >>> print(text[:100])
        """
        try:
            if not os.path.exists(file_path):
                raise FileNotFoundError(f"PDF file not found: {file_path}")
            
            logger.info(f"Loading PDF: {file_path}")
            
            file_size = os.path.getsize(file_path)
            if file_size > MAX_FILE_SIZE:
                raise ValueError(f"File too large: {file_size / 1024 / 1024:.1f} MB")
            
            # Open and read PDF
            text_content = []
            page_count = 0
            
            with open(file_path, 'rb') as pdf_file:
                pdf_reader = PyPDF2.PdfReader(pdf_file)
                page_count = len(pdf_reader.pages)
                
                logger.info(f"PDF has {page_count} pages")
                
                for page_num, page in enumerate(pdf_reader.pages, 1):
                    try:
                        page_text = page.extract_text()
                        if page_text:
                            text_content.append(page_text)
                    except Exception as e:
                        logger.warning(f"Failed to extract page {page_num}: {str(e)}")
                        continue
            
            full_text = "\n\n".join(text_content)
            
            logger.info(f"Successfully extracted {len(full_text)} characters from {page_count} pages")
            return full_text
            
        except Exception as e:
            logger.error(f"Error loading PDF: {str(e)}")
            raise
    
    def load_with_metadata(self, file_path: str) -> Tuple[str, Dict]:
        """
        Load PDF with metadata.
        
        Returns:
            Tuple of (text, metadata)
        """
        text = self.load(file_path)
        metadata = {
            'source': os.path.basename(file_path),
            'file_type': 'PDF',
            'file_size': os.path.getsize(file_path),
            'pages': len(text.split('\n\n'))  # Approximate
        }
        return text, metadata


# ============================================================================
# TEXT LOADER
# ============================================================================

class TextLoader(BaseFileLoader):
    """
    Loader for plain text files (.txt, .md).
    
    Reads and processes text-based files.
    
    Attributes:
        encoding: Text file encoding (default: utf-8)
    """
    
    def __init__(self, encoding: str = 'utf-8'):
        """
        Initialize Text Loader.
        
        Args:
            encoding: File encoding to use
        """
        self.encoding = encoding
    
    def supports_format(self, file_path: str) -> bool:
        """Check if file is text-based."""
        return file_path.lower().endswith(('.txt', '.md'))
    
    def load(self, file_path: str) -> str:
        """
        Load text from file.
        
        Args:
            file_path: Path to text file
            
        Returns:
            File contents as string
            
        Example:
            >>> loader = TextLoader()
            >>> text = loader.load("document.txt")
        """
        try:
            if not os.path.exists(file_path):
                raise FileNotFoundError(f"Text file not found: {file_path}")
            
            logger.info(f"Loading text file: {file_path}")
            
            file_size = os.path.getsize(file_path)
            if file_size > MAX_FILE_SIZE:
                raise ValueError(f"File too large: {file_size / 1024 / 1024:.1f} MB")
            
            # Read file with specified encoding
            with open(file_path, 'r', encoding=self.encoding) as file:
                text = file.read()
            
            logger.info(f"Successfully loaded {len(text)} characters")
            return text
            
        except UnicodeDecodeError:
            logger.warning(f"UTF-8 decode failed, trying latin-1")
            with open(file_path, 'r', encoding='latin-1') as file:
                text = file.read()
            return text
            
        except Exception as e:
            logger.error(f"Error loading text file: {str(e)}")
            raise
    
    def load_with_metadata(self, file_path: str) -> Tuple[str, Dict]:
        """Load text file with metadata."""
        text = self.load(file_path)
        metadata = {
            'source': os.path.basename(file_path),
            'file_type': 'Text',
            'file_size': os.path.getsize(file_path),
            'lines': len(text.split('\n'))
        }
        return text, metadata


# ============================================================================
# WORD DOCUMENT LOADER
# ============================================================================

class DocxLoader(BaseFileLoader):
    """
    Loader for Word documents (.docx).
    
    Extracts text from DOCX files.
    """
    
    def supports_format(self, file_path: str) -> bool:
        """Check if file is DOCX."""
        return file_path.lower().endswith('.docx')
    
    def load(self, file_path: str) -> str:
        """
        Extract text from DOCX file.
        
        Args:
            file_path: Path to DOCX file
            
        Returns:
            Extracted text
            
        Example:
            >>> loader = DocxLoader()
            >>> text = loader.load("document.docx")
        """
        try:
            if not os.path.exists(file_path):
                raise FileNotFoundError(f"DOCX file not found: {file_path}")
            
            logger.info(f"Loading DOCX: {file_path}")
            
            # Try to import python-docx
            try:
                from docx import Document as DocxDocument
            except ImportError:
                logger.error("python-docx not installed. Install with: pip install python-docx")
                raise
            
            # Load DOCX
            doc = DocxDocument(file_path)
            
            # Extract text from all paragraphs
            paragraphs = []
            for para in doc.paragraphs:
                if para.text.strip():
                    paragraphs.append(para.text)
            
            full_text = "\n".join(paragraphs)
            
            logger.info(f"Successfully extracted {len(full_text)} characters from {len(doc.paragraphs)} paragraphs")
            return full_text
            
        except Exception as e:
            logger.error(f"Error loading DOCX: {str(e)}")
            raise
    
    def load_with_metadata(self, file_path: str) -> Tuple[str, Dict]:
        """Load DOCX with metadata."""
        text = self.load(file_path)
        metadata = {
            'source': os.path.basename(file_path),
            'file_type': 'DOCX',
            'file_size': os.path.getsize(file_path),
        }
        return text, metadata


# ============================================================================
# DOCUMENT LOADER CLASS (MAIN)
# ============================================================================

class DocumentLoader:
    """
    Main document loader that handles multiple file formats.
    
    This class:
    1. Detects file type
    2. Uses appropriate loader
    3. Extracts and cleans text
    4. Splits into chunks
    5. Attaches metadata
    
    Attributes:
        chunk_size: Size of text chunks
        chunk_overlap: Overlap between chunks
        loaders: Dictionary of file type loaders
        text_splitter: RecursiveCharacterTextSplitter instance
    """
    
    def __init__(
        self,
        chunk_size: int = 500,
        chunk_overlap: int = 100,
        encoding: str = 'utf-8'
    ):
        """
        Initialize Document Loader.
        
        Args:
            chunk_size: Characters per chunk
            chunk_overlap: Character overlap between chunks
            encoding: Text file encoding
        """
        self.chunk_size = chunk_size
        self.chunk_overlap = chunk_overlap
        self.encoding = encoding
        
        # Initialize file loaders
        self.loaders = {
            '.pdf': PDFLoader(),
            '.txt': TextLoader(encoding),
            '.md': TextLoader(encoding),
            '.docx': DocxLoader(),
        }
        
        # Initialize text splitter
        self.text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=chunk_size,
            chunk_overlap=chunk_overlap,
            separators=["\n\n", "\n", ". ", " ", ""]
        )
        
        logger.info(f"DocumentLoader initialized (chunk_size={chunk_size}, overlap={chunk_overlap})")
    
    # ========================================================================
    # FILE LOADING
    # ========================================================================
    
    def load_file(
        self,
        file_path: str,
        return_metadata: bool = False
    ) -> List[str]:
        """
        Load and chunk a single file.
        
        Args:
            file_path: Path to document file
            return_metadata: If True, return tuples of (chunk, metadata)
            
        Returns:
            List of text chunks (or tuples if return_metadata=True)
            
        Example:
            >>> loader = DocumentLoader()
            >>> chunks = loader.load_file("document.pdf")
            >>> print(f"Got {len(chunks)} chunks")
            >>> print(chunks[0][:100])
        """
        try:
            logger.info(f"Loading file: {file_path}")
            
            # Check file exists
            if not os.path.exists(file_path):
                raise FileNotFoundError(f"File not found: {file_path}")
            
            # Get file extension
            file_ext = os.path.splitext(file_path)[1].lower()
            
            # Check if supported
            if file_ext not in self.loaders:
                supported = ", ".join(self.loaders.keys())
                raise ValueError(f"Unsupported file type: {file_ext}. Supported: {supported}")
            
            # Load file with appropriate loader
            loader = self.loaders[file_ext]
            text = loader.load(file_path)
            
            if not text or not text.strip():
                logger.warning("Extracted text is empty")
                return []
            
            # Clean text
            text = self._clean_text(text)
            
            # Split into chunks
            chunks = self.text_splitter.split_text(text)
            
            logger.info(f"Split into {len(chunks)} chunks")
            
            if return_metadata:
                _, metadata = loader.load_with_metadata(file_path)
                return [(chunk, metadata) for chunk in chunks]
            
            return chunks
            
        except Exception as e:
            logger.error(f"Error loading file: {str(e)}")
            raise
    
    def load_directory(
        self,
        directory_path: str,
        recursive: bool = True
    ) -> List[str]:
        """
        Load and chunk all documents in a directory.
        
        Args:
            directory_path: Path to directory
            recursive: Whether to search subdirectories
            
        Returns:
            List of all chunks from all files
            
        Example:
            >>> loader = DocumentLoader()
            >>> chunks = loader.load_directory("./documents/")
            >>> print(f"Got {len(chunks)} chunks from multiple files")
        """
        try:
            logger.info(f"Loading directory: {directory_path}")
            
            if not os.path.isdir(directory_path):
                raise ValueError(f"Not a directory: {directory_path}")
            
            all_chunks = []
            
            # Get file pattern
            if recursive:
                files = Path(directory_path).rglob('*')
            else:
                files = Path(directory_path).glob('*')
            
            # Process each file
            file_count = 0
            for file_path in files:
                if file_path.is_file():
                    file_ext = file_path.suffix.lower()
                    if file_ext in self.loaders:
                        try:
                            chunks = self.load_file(str(file_path))
                            all_chunks.extend(chunks)
                            file_count += 1
                        except Exception as e:
                            logger.warning(f"Failed to load {file_path}: {str(e)}")
            
            logger.info(f"Loaded {file_count} files, total {len(all_chunks)} chunks")
            return all_chunks
            
        except Exception as e:
            logger.error(f"Error loading directory: {str(e)}")
            raise
    
    # ========================================================================
    # TEXT PROCESSING
    # ========================================================================
    
    def _clean_text(self, text: str) -> str:
        """
        Clean and normalize text.
        
        Args:
            text: Raw text to clean
            
        Returns:
            Cleaned text
        """
        # Remove extra whitespace
        text = " ".join(text.split())
        
        # Remove special characters that break processing
        text = text.replace('\x00', '')  # Null bytes
        text = text.replace('\ufffd', '')  # Replacement character
        
        # Remove excessive newlines
        while '\n\n\n' in text:
            text = text.replace('\n\n\n', '\n\n')
        
        return text
    
    def chunk_text(
        self,
        text: str,
        chunk_size: Optional[int] = None,
        chunk_overlap: Optional[int] = None
    ) -> List[str]:
        """
        Chunk raw text without loading from file.
        
        Args:
            text: Text to chunk
            chunk_size: Override default chunk size
            chunk_overlap: Override default overlap
            
        Returns:
            List of text chunks
            
        Example:
            >>> loader = DocumentLoader()
            >>> chunks = loader.chunk_text(
            ...     "Long text...",
            ...     chunk_size=200,
            ...     chunk_overlap=50
            ... )
        """
        try:
            # Use custom splitter if different size requested
            if chunk_size or chunk_overlap:
                splitter = RecursiveCharacterTextSplitter(
                    chunk_size=chunk_size or self.chunk_size,
                    chunk_overlap=chunk_overlap or self.chunk_overlap
                )
            else:
                splitter = self.text_splitter
            
            chunks = splitter.split_text(text)
            logger.info(f"Chunked text into {len(chunks)} pieces")
            return chunks
            
        except Exception as e:
            logger.error(f"Error chunking text: {str(e)}")
            raise
    
    def get_file_stats(self, file_path: str) -> Dict:
        """
        Get statistics about a file without fully loading it.
        
        Args:
            file_path: Path to file
            
        Returns:
            Dictionary with file statistics
            
        Example:
            >>> stats = loader.get_file_stats("document.pdf")
            >>> print(f"File size: {stats['file_size']} bytes")
        """
        try:
            if not os.path.exists(file_path):
                raise FileNotFoundError(f"File not found: {file_path}")
            
            file_stats = {
                'file_name': os.path.basename(file_path),
                'file_type': os.path.splitext(file_path)[1],
                'file_size': os.path.getsize(file_path),
                'file_size_mb': os.path.getsize(file_path) / (1024 * 1024),
                'modified_time': os.path.getmtime(file_path),
            }
            
            return file_stats
            
        except Exception as e:
            logger.error(f"Error getting file stats: {str(e)}")
            raise
    
    # ========================================================================
    # BATCH PROCESSING
    # ========================================================================
    
    def load_files(self, file_paths: List[str]) -> List[str]:
        """
        Load multiple files and return all chunks.
        
        Args:
            file_paths: List of file paths
            
        Returns:
            Combined list of all chunks
            
        Example:
            >>> loader = DocumentLoader()
            >>> chunks = loader.load_files([
            ...     "doc1.pdf",
            ...     "doc2.txt",
            ...     "doc3.docx"
            ... ])
        """
        all_chunks = []
        
        for file_path in file_paths:
            try:
                chunks = self.load_file(file_path)
                all_chunks.extend(chunks)
                logger.info(f"Loaded {len(chunks)} chunks from {os.path.basename(file_path)}")
            except Exception as e:
                logger.warning(f"Failed to load {file_path}: {str(e)}")
                continue
        
        logger.info(f"Total loaded: {len(all_chunks)} chunks from {len(file_paths)} files")
        return all_chunks


# ============================================================================
# HELPER FUNCTIONS
# ============================================================================

def is_supported_format(file_path: str) -> bool:
    """
    Check if file format is supported.
    
    Args:
        file_path: Path to file
        
    Returns:
        bool: True if supported
        
    Example:
        >>> is_supported_format("document.pdf")  # True
        >>> is_supported_format("image.jpg")      # False
    """
    file_ext = os.path.splitext(file_path)[1].lower()
    return file_ext in SUPPORTED_FORMATS


def get_supported_formats() -> Dict[str, str]:
    """
    Get all supported file formats and descriptions.
    
    Returns:
        Dictionary of formats and descriptions
    """
    return SUPPORTED_FORMATS.copy()


def estimate_chunks(file_size_mb: float, chunk_size: int = 500) -> int:
    """
    Estimate how many chunks a file will produce.
    
    Args:
        file_size_mb: File size in megabytes
        chunk_size: Characters per chunk
        
    Returns:
        Estimated number of chunks
        
    Example:
        >>> chunks = estimate_chunks(5.0)  # 5 MB file
        >>> print(f"Approximately {chunks} chunks")
    """
    # Rough estimate: 1 MB ≈ 1 million characters
    # Each chunk ≈ chunk_size characters
    return int((file_size_mb * 1_000_000) / chunk_size)
